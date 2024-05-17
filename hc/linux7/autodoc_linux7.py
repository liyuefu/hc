##coding=utf-8
#!/usr/bin/env python
from docx import *
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches #支持修改图片大小的库
from docx.shared import RGBColor#设置字体
import codecs
from codecs import open
import re
import cx_Oracle
import subprocess
import commands
from datetime import date
#from configobj import ConfigObj
import ConfigParser
import sys
import os
reload(sys)
import logging

PWD=os.getcwd()
logging.basicConfig(level=logging.DEBUG,format='%(asctime)s %(name)-12s %(levelname)-8s %(message)s',datefmt='%m-%d %H:%M',filename='./autodoc_linux7.log',filemode='w')
logging.info('starting autodoc_linux7.py now...')
sys.setdefaultencoding("utf-8")
sqlfilename=sys.argv[1]
oraclehome=sys.argv[2]
dbver2=sys.argv[3]
ip_address=sys.argv[4]
jpgfilepath=PWD+"/hcscripts/"
"""
Desc: check database status and create word document.
Env: RH6. Python2.6. Oracle 11.2.0.4
package needed :docx, cx_Oracle
files list: 1. autodoc.sql (database scripts). 2. autodoc.py(this script)

what this scirpt do
1. run sql script to collect database info.
2. run linux command to collect os info.
3. write information to doc.

to be finished by yourself:
4. add awr report performance info to doc.
5. analyze alert.log, awr report and other sql infomation to provide suggestions.
6. recreate "table of contents"
7. change title/client name, etc.

update record:
create: 2020.11.16 V1.3
update: 2020.11.17 V1.4.  
        1.correct fetch NLS_CHARACTER etc sequence error.
        2.save name with hostname_dbname_date
update: 2020.11.19. rewrite for linux7
update: 2020.11.26. add try except for subprocess linux command
update: 2020.11.27. add font change.
update: 2021.02.01  change linuxoutput format to size 10
update: 2021.06.20. change table first row color
update: 2021.09.02 v20210902.
        fixed lspatch exception catch bug.
update: 2022.06.13. v0220613 red config info from config.ini.
update: 2022.06.18. read ip address,disaster recovery from config.ini, 简化巡检总结
update: 2022.12.21. fixed 19c v$pwfile_users 超过最大12个字段的问题. 调整为最大20个字段.增加debug
"""

##########################read config.ini##############################

PWD=os.getcwd()
config = ConfigParser.ConfigParser()
config.read(PWD+"/config.ini")

try:
  client_name = config.get('client','client_name')
  client_app_name = config.get('client','client_app_name')
  client_app_db_name = config.get('client','client_app_db_name')
  client_dba_name = config.get('client','client_dba_name')
  disaster_recovery = config.get('client','disaster_recovery')
  lidao_tech_manager = config.get('lidao','lidao_tech_manager')
  lidao_sales_name = config.get('lidao','lidao_sales_name')
  lidao_engineer_name = config.get('lidao','lidao_engineer_name')
  lidao_header=config.get('lidao','lidao_header')
  lidao_copyright=config.get('lidao','lidao_copyright')
  lidao_name=config.get('lidao','lidao_name')
  lidao_address=config.get('lidao','lidao_address')
  lidao_tel=config.get('lidao','lidao_tel')
except Exception as reason:
  logging.error(reason)
  print("read config.ini failed")


client_name = str(client_name).encode("utf-8").decode("utf-8")
client_app_name = str(client_app_name).encode("utf-8").decode("utf-8")
client_app_db_name = str(client_app_db_name).encode("utf-8").decode("utf-8")
client_dba_name = str(client_dba_name).encode("utf-8").decode("utf-8")
disaster_recovery = str(disaster_recovery).encode("utf-8").decode("utf-8")
lidao_tech_manager = str(lidao_tech_manager).encode("utf-8").decode("utf-8")
lidao_sales_name = str(lidao_sales_name).encode("utf-8").decode("utf-8")
lidao_engineer_name = str(lidao_engineer_name).encode("utf-8").decode("utf-8")
lidao_header = str(lidao_header).encode("utf-8").decode("utf-8")
lidao_copyright = str(lidao_copyright).encode("utf-8").decode("utf-8")
lidao_name = str(lidao_name).encode("utf-8").decode("utf-8")
lidao_address = str(lidao_address).encode("utf-8").decode("utf-8")
lidao_tel = str(lidao_tel).encode("utf-8").decode("utf-8")

# GETIP=PWD+"/hcscripts/getip.sh"
# try:
#   res,ip_address=commands.getstatusoutput(GETIP)
#   logging.info(ip_address)
# except:
#   logging.error("getip fail")
########################get patch info #################################
opatch = oraclehome+'/OPatch/opatch'
opatchtable=[[ur"数据库补丁信息"]]
try:
    if (dbver2 > "11.2.0.3.0"):
        ps = subprocess.Popen([opatch,' lspatches'],stdout=subprocess.PIPE)
    else:
        ps = subprocess.Popen([opatch,' lsinventory'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]
    for line in output.splitlines():
        opatchtable.append([line])
except Exception as reason:
    logging.error(reason)
    print ("opatch  failed")

try:
    subprocess.check_output([opatch,' lspatches'])
except subprocess.CalledProcessError as reason:
    logging.info(reason)
    logging.info("opatch versio is low")
    print("OPatch version is low,no support lspatches,try lsinventory")
    ps = subprocess.Popen([opatch,' lsinventory'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]
else:
    print("----------------opatch lspatches OK.-----------")
    ps = subprocess.Popen([opatch,' lspatches'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]

for line in output.splitlines():
    opatchtable.append([line])
#****************************This part create word doc.************************

#create word document
#put ur before Chinese  str
# Create our properties, contenttypes, and other support files
print(ip_address)
filename = client_name + "_" + client_app_name 
filename_end="_" + client_app_db_name + "健康检查报告"
#date of today
today=str(date.today())
filename_end=filename_end+"_"+today;

#define Normal font
document = Document()
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal'].font.size = Pt(12)
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

#define linux command output format
styles = document.styles
new_heading_style = styles.add_style('linuxoutput', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Normal']
font = new_heading_style.font
font.name = ur'宋体'
font.size = Pt(8)

#define New Heading1
styles = document.styles
new_heading_style = styles.add_style('New Heading1', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Heading 1']
font = new_heading_style.font
font.name = ur'宋体'
font.size = Pt(18)
document.styles['New Heading1'].font.name = u'宋体'
document.styles['New Heading1']._element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')

#define New Heading2
styles = document.styles
new_heading_style = styles.add_style('New Heading2', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Heading 2']
font = new_heading_style.font
font.name = ur'宋体'
font.size = Pt(16)
document.styles['New Heading2'].font.name = u'宋体'
document.styles['New Heading2']._element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')

#define New Heading4
styles = document.styles
new_heading_style = styles.add_style('New Heading4', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Heading 3']
font = new_heading_style.font
font.name = ur'宋体'
font.size = Pt(12)
document.styles['New Heading4'].font.name = u'宋体'
document.styles['New Heading4']._element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')

#define New Heading5
styles = document.styles
new_heading_style = styles.add_style('New Heading5', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Heading 5']
font = new_heading_style.font
font.name = ur'黑体'
font.color.rgb = RGBColor(0,0,255)  #设置颜色为兰色
font.size = Pt(13)
document.styles['New Heading5'].font.name = u'黑体'
document.styles['New Heading5']._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')

############head ######
#页眉格式
styles = document.styles
style = styles.add_style("NewHeader", WD_STYLE_TYPE.PARAGRAPH)
style.base_style = styles["Normal"]
tab_stops = style.paragraph_format.tab_stops
tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
document.styles['NewHeader'].font.size = Pt(7)
document.styles['NewHeader'].font.underline = True

#页眉
section = document.sections[0]
header = section.header
headparagraph = header.paragraphs[0]
#页眉加图片

#headparagraph.text = ur"利道全国7×24技术服务热线:    15821486451    021-53083780                             公司网站：http://www.leaderit.com"
headparagraph.text=lidao_header
headparagraph.style = document.styles["NewHeader"]
#r = headparagraph.add_run()
#r.add_picture(jpgfilepath+'leaderit2.jpg', width=Inches(1))

document.add_paragraph("\n")
document.add_picture(jpgfilepath+'leaderit.jpg', width=Inches(4)) #添加图片，并设置宽度

document.add_paragraph("\n")
document.add_paragraph("\n")


p=document.add_paragraph()
#run=p.add_run(ur'客户名称')
run=p.add_run(client_name)
run.font.name=(ur'黑体')
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
run.font.size = Pt(24)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p=document.add_paragraph()
#run=p.add_run(ur'应用系统')
app_db=client_app_name + client_app_db_name
run=p.add_run(app_db)
run.font.name=(ur'黑体')
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
run.font.size = Pt(24)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p=document.add_paragraph()
#run=p.add_run(ur'数据库健康检查报告')
run=p.add_run(ur'健康检查报告')
run.font.name=(ur'黑体')
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
run.font.size = Pt(24)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p=document.add_paragraph()
#run=p.add_run(ur'Copyright © leaderit software Corporation  保留所有权利')
run=p.add_run(lidao_copyright)
run.font.name=(ur'黑体')
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
run.font.size = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER


emptyline=4
while emptyline>0:
    document.add_paragraph(("\n"))
    emptyline -=1;

#document.add_paragraph(ur'上海利道软件技术有限公司' )
#document.add_paragraph(ur'上海市北京东路 668号科技京城东27楼A3B3座 ' )
#document.add_paragraph(ur'电话：（86-21）53083780 传真：（86-21）53083787' )
document.add_paragraph(lidao_name )
document.add_paragraph(lidao_address)
document.add_paragraph(lidao_tel)

document.add_page_break()


#version control page
document.add_paragraph(ur'文档控制', style='New Heading1')
document.add_paragraph(ur'修改记录', style='New Heading2')

records =[ 
    [ur'日期',ur'作者',ur'版本',ur'修改记录'],
    [today, lidao_engineer_name, ur'V1.0',ur'创建文档']
]

first_time=[today,lidao_engineer_name,client_dba_name]
table = document.add_table(0, len(records[0]))
table.style='Table Grid'
for row in records:
    cells=table.add_row().cells
    col_no = 0
    for col in row:
        cells[col_no].text=col
        col_no += 1
#set backcolor of first cell begin
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_2 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_3 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_4 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_5 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_6 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_7 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_8 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_9 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_a = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_list=[shading_elm_1,shading_elm_2,shading_elm_3,shading_elm_4,shading_elm_5,shading_elm_6,shading_elm_7,shading_elm_8,shading_elm_9,shading_elm_a]
col_no=0
while col_no<len(records[0]):
  table.rows[0].cells[col_no]._tc.get_or_add_tcPr().append(shading_list[col_no])
  col_no += 1
#set backcolor of first cell end


document.add_paragraph("\n")
document.add_paragraph("\n")

document.add_paragraph(ur'审阅记录', style='New Heading2')

records=[[ur"审   阅   人",ur" 职                      位"],[lidao_tech_manager,ur"技术总监"]]
table = document.add_table(0, len(records[0]))
table.style='Table Grid'
for row in records:
    cells=table.add_row().cells
    col_no = 0
    for col in row:
        cells[col_no].text=col
        col_no += 1

#set backcolor of first cell begin
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_2 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_3 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_4 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_5 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_6 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_7 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_8 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_9 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_a = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_list=[shading_elm_1,shading_elm_2,shading_elm_3,shading_elm_4,shading_elm_5,shading_elm_6,shading_elm_7,shading_elm_8,shading_elm_9,shading_elm_a]
col_no=0
while col_no<len(records[0]):
  table.rows[0].cells[col_no]._tc.get_or_add_tcPr().append(shading_list[col_no])
  col_no += 1
#set backcolor of first cell end

document.add_paragraph("\n")
document.add_paragraph("\n")


document.add_paragraph(ur'分发记录', style='New Heading2')
#records=[[ur"拷贝编号",ur" 姓    名",ur"单             位"],["1",client_dba_name,client_name],["2",ur"利道技术部",ur"利道软件"]]
records=[[ur"拷贝编号",ur" 姓    名",ur"单            位"],["1",client_dba_name,client_name],["2",ur"技术部",lidao_name], ["3",lidao_sales_name,
lidao_name]]
table = document.add_table(0, len(records[0]))
table.style='Table Grid'
for row in records:
    cells=table.add_row().cells
    col_no = 0
    for col in row:
        cells[col_no].text=col
        col_no += 1

#set backcolor of first cell begin
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_2 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_3 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_4 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_5 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_6 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_7 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_8 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_9 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_a = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_list=[shading_elm_1,shading_elm_2,shading_elm_3,shading_elm_4,shading_elm_5,shading_elm_6,shading_elm_7,shading_elm_8,shading_elm_9,shading_elm_a]
col_no=0
while col_no<len(records[0]):
  table.rows[0].cells[col_no]._tc.get_or_add_tcPr().append(shading_list[col_no])
  col_no += 1
#set backcolor of first cell end
#table of contents
document.add_page_break()

#get db basic info
#****************************This part create db basic info table .************************
#connect to oracle  with / as sysdba

#con = cx_Oracle.connect('sys','manager',mode=cx_Oracle.SYSDBA)
con = cx_Oracle.connect("/", mode = cx_Oracle.SYSDBA )
cur = con.cursor()

db_info=[[ur"参数",ur"参数值"]]
#db_info.append([ur'应用名称',client_app_name])
db_info.append([ur'可用性要求','M-F 5x24'])
db_info.append([ur'容灾方案',disaster_recovery])

#controlfile 
sql="select name from v$controlfile"
cur.execute(sql)
result = cur.fetchall()
allrecords = ''
for onerow in result:
    allrecords = allrecords +  str(onerow[0]) + ','

ctlfile = [ur'控制文件',allrecords]

db_info.append(ctlfile)


#cluster or not
sql="select value from v$parameter where name = 'cluster_database'"
cur.execute(sql)
col1 = cur.fetchone()
cluster = [ur'是否RAC',str(col1[0])]

db_info.append(cluster)


#dataguard or not
feature_boolean = cur.var(int)
aux_count = cur.var(int)
feature_info=cur.var(str)
cur.callproc('dbms_feature_data_guard', [feature_boolean, aux_count, feature_info])
if feature_boolean.getvalue()>0:
  standby = [ur'是否启用DG',ur'是']
else:
  standby = [ur'是否启用DG',ur'否']
db_info.append(standby)


#log_mode,flashback_on
sql="select log_mode,flashback_on from v$database"
cur.execute(sql)
col1,col2 = cur.fetchone()
log_mode = [ur'归档模式',str(col1)]
flashback_on=[ur'是否开启FLASHBACK',str(col2)]
db_info.append(log_mode)
db_info.append(flashback_on)

sql="select * from V$version"
cur.execute(sql)
result = cur.fetchone()
version = [ur'数据库版本',str(result[0])]

db_info.append(version)

sql="select name from v$database"
cur.execute(sql)
col1 = cur.fetchone()
dbname = [ur'数据库名',str(col1[0])]
dbnamestr=str(col1[0]);

db_info.append(dbname)

sql=" select count(*) from v$log"
cur.execute(sql)
col1=cur.fetchone()
logcnt=[ur'日志组数',str(col1[0])]

db_info.append(logcnt)

sql = "select distinct cnt from ( select group#,count(*) cnt from v$logfile group by group#) where rownum=1"
cur.execute(sql)
col1=cur.fetchone()
members=[ur'日志组成员数',str(col1[0])]

db_info.append(members)

sql="select  bytes/1024/1024 M from v$log where rownum=1"
cur.execute(sql)
col1=cur.fetchone()
logsize=[ur'日志文件(M)',str(col1[0])]

db_info.append(logsize)


sql="select  c.sum3 DMP_G ,a.sum1 RMAN_G,b.sum2 DATA_G from (SELECT ceil(SUM(BYTES)/1024/1024/1024) sum1 FROM DBA_segments ) a,(select ceil(sum(bytes)/1024/1024/1024) sum2 from v$datafile) b,(select ceil(sum(bytes)/1024/1024/1024) sum3 from dba_segments where segment_type not like 'INDEX%' and segment_type not in('ROLLBACK','CACHE','LOBINDEX','TYPE2 UNDO')) c "

cur.execute(sql)
col1,col2,col3=cur.fetchone()
dmp=str(col1)
rman=str(col2)
data=str(col3)
dmpsize=[ur'数据泵文件大小(G)',dmp]
rmansize=[ur'RMAN备份大小(G)',rman]
datasize=[ur'数据文件大小(G)',data]
db_info.append(dmpsize)
db_info.append(rmansize)
db_info.append(datasize)

sql="select value$ from sys.props$ where name in('NLS_TERRITORY','NLS_LANGUAGE','NLS_CHARACTERSET','NLS_NCHAR_CHARACTERSET') order by name"
cur.execute(sql)
col1=cur.fetchone()
col2=cur.fetchone()
col3=cur.fetchone()
col4=cur.fetchone()
nls_char=['NLS_CHARACTERSET',str(col1[0])]
nls_lang=['NLS_LANGUAGE',str(col2[0])]
nls_nchar=['NLS_NCHAR_CHARACTERSET',str(col3[0])]
nls_terr=['NLS_TERRITORY',str(col4[0])]

db_info.append(nls_terr)
db_info.append(nls_lang)
db_info.append(nls_char)
db_info.append(nls_nchar)

#sql="select value from dba_hist_osstat where stat_name = 'NUM_CPUS' and rownum=1"
sql = "select value from v$parameter where name = 'cpu_count' and rownum=1"

cur.execute(sql)
col1=cur.fetchone()
try:
    num_cpus=[ur'CPU数量',str(col1[0])]
    db_info.append(num_cpus)
except Exception  as reason:
    logging.error(reason)
    logging.error(sql)
    print("sql error")
    print(sql)

sql="select value/1024/1024/1024 from dba_hist_osstat where stat_name = 'PHYSICAL_MEMORY_BYTES' and rownum=1"
cur.execute(sql)
row=cur.fetchone()
try:
    col1=round(float(str(row[0])),1)
    ram_g=[ur'内存 G',str(col1)]
    db_info.append(ram_g)
except Exception as reason:
    logging.error(reason)
    logging.error(sql)
    print("sql error")
    print(sql)

#os=["OS Ver","Please input OS ver/platform/cpu/RAM here       "]
#db_info.append(os)

try:
    ps = subprocess.Popen(['cat', '/etc/redhat-release'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]
except Exception  as reason:
    logging.error(reason)
    print("cat redhat releae fail")
osver=""
for line in output.splitlines():
    osver=osver+str(line) 
ver=[ur'操作系统版本',osver]
db_info.append(ver)

ip=[ur"IP 地址",ip_address]
db_info.append(ip)

# document.add_page_break()
document.add_paragraph(ur'一. 数据库基本信息', style='New Heading1')

table = document.add_table(0, len(db_info[0]))
table.style='Table Grid'
for row in db_info:
    col_no=0
    cells = table.add_row().cells
    for col in row:
        cells[col_no].text=col
        col_no += 1

#set backcolor of first cell begin
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_2 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_3 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_4 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_5 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_6 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_7 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_8 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_9 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_a = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_list=[shading_elm_1,shading_elm_2,shading_elm_3,shading_elm_4,shading_elm_5,shading_elm_6,shading_elm_7,shading_elm_8,shading_elm_9,shading_elm_a]
col_no=0
while col_no<len(db_info[0]):
  table.rows[0].cells[col_no]._tc.get_or_add_tcPr().append(shading_list[col_no])
  col_no += 1
#set backcolor of first cell end
document.add_page_break()


#****************************This part create db detail .************************
#service main content
document.add_paragraph(ur'二. 数据库巡检记录', style='New Heading1')

records=[[ur"日期",ur"利道工程师",ur"客户工程师"]]

#history service record
sections=config.sections()
pattern='^date'
for section in sections:
  match = re.match(pattern,section)
  if match:
    his_date = config.get(section,'date')
    first_time=[his_date,lidao_engineer_name,client_dba_name]
    records.append(first_time)

first_time=[today,lidao_engineer_name,client_dba_name]
records.append(first_time)


table = document.add_table(0, len(records[0]))
table.style='Table Grid'
for row in records:
    col_no=0
    cells = table.add_row().cells
    for col in row:
        cells[col_no].text=col
        col_no += 1

#set backcolor of first cell begin
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_2 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_3 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_4 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_5 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_6 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_7 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_8 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_9 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_a = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_list=[shading_elm_1,shading_elm_2,shading_elm_3,shading_elm_4,shading_elm_5,shading_elm_6,shading_elm_7,shading_elm_8,shading_elm_9,shading_elm_a]
col_no=0
while col_no<len(records[0]):
  table.rows[0].cells[col_no]._tc.get_or_add_tcPr().append(shading_list[col_no])
  col_no += 1
#set backcolor of first cell end
document.add_paragraph("\n")
document.add_paragraph("\n")

document.add_paragraph(ur'三. 数据增长趋势', style='New Heading1')
#records=[[ur'日期',"DMP(G)","RMAN(G)","DATA(G)"],[today,dmp,rman,data]]
records=[[ur'日期',"DMP(G)","RMAN(G)","DATA(G)"]]
table = document.add_table(0, len(records[0]))

#read history data from config.ini
sections=config.sections()
pattern='^date'
for section in sections:
  match = re.match(pattern,section)
  if match:
    his_date = config.get(section,'date')
    his_dmp = config.get(section,'dmp')
    his_rman = config.get(section,'rman')
    his_data = config.get(section,'data')
    datagrow=[his_date,his_dmp,his_rman,his_data]
    records.append(datagrow)

datagrow=[today,dmp,rman,data]
records.append(datagrow)

#write check information into config.ini
date='date-'+today
try: 
    config.add_section(date)
    config.set(date,'date',today)
    config.set(date,'dmp',dmp)
    config.set(date,'rman',rman)
    config.set(date,'data',data)
    f=open((PWD+"/config.ini"),'w')
    config.write(f)
except Exception  as reason:
    logging.error(reason)
    print("write check history record to config.ini failed.")
    logging.error("write check history record to config.ini failed.")
    

table.style='Table Grid'
for row in records:
    col_no=0
    cells = table.add_row().cells
    for col in row:
        cells[col_no].text=col
        col_no += 1

#set backcolor of first cell begin
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_2 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_3 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_4 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_5 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_6 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_7 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_8 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_9 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_a = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_list=[shading_elm_1,shading_elm_2,shading_elm_3,shading_elm_4,shading_elm_5,shading_elm_6,shading_elm_7,shading_elm_8,shading_elm_9,shading_elm_a]
col_no=0
while col_no<len(records[0]):
  table.rows[0].cells[col_no]._tc.get_or_add_tcPr().append(shading_list[col_no])
  col_no += 1
#set backcolor of first cell end
document.add_page_break()


document.add_paragraph(ur'四. 数据库逐项检查', style='New Heading1')


#sqllist: store the sql command from sqlfile.
#sqldesclist: store the sql  desc from sqlfile.
sqllist=ur""
sqldesclist=ur""

#******************this part connect to db and fetch data and write to word doc*********************
#read from sqlfile and store command/desc in sqllist/sqldesclist
tmpstr=""
with open(sqlfilename) as f:
    for line in f:
        if not(re.match("prompt",line) or re.match("--",line) or  re.match("exit",line) or
            re.match("col",line) or re.match("alter",line) or re.match("set",line) or re.match("spool",line)):
            sqllist = sqllist + line
        elif (re.match("prompt",line)):
            if (re.match("prompt \'[0-9]*,",line)):
                sqldesclist = sqldesclist + line[6:] +";";                
                
#seperate sql script with ;and store in sqlcommand list
#seperate sqldesc with ; and store in sqldesc list
sqlcommand = sqllist.split(';')
sqldesc = sqldesclist.split(';')



#run every sql command in sqlcommand list.
#write the sql output to doc in table.

k=0
while (k < len(sqlcommand)-1 and  k < len(sqldesc)-1):
    logging.info(sql)
    sql = sqlcommand[k]
    sqltext=sqldesc[k]
    cur = con.cursor()
    try:
        cur.execute(sql)
    except Exception  as reason:
        logging.error(reason)
        logging.error(sql)
        print sql
        print "Execute error ."  
        onerow="    "
    try:
        result = cur.fetchall()
    except Exception as reason:
        logging.error(reason)
        logging.error(sql)
        print "execute fail"
        print sql
    col_names = [row[0] for row in cur.description or []]
    numcols=len(col_names)
    numrows=len(result)
    allrow=[col_names]
    i=0
    while i< numrows:
        j=0
        newrow=[]
        while j<numcols:
            col=str(result[i][j])
            newrow.append(col)
            j=j+1
        allrow.append(newrow)
        newrow=[]
        i=i+1
#write title for each sql command.
    if (re.search("MAX OBJECT ID",sqltext)):
        document.add_paragraph(ur"(一). 检查潜在的风险",style='New Heading2')
    elif (re.search("DB_NAME LOGMODE",sqltext)):
        document.add_paragraph(ur"(二).数据库一览",style='New Heading2' )
    elif (re.search("DB CONFIGURATION",sqltext)):
        document.add_paragraph(ur"(三).数据库配置",style='New Heading2')
    elif (re.search("CONTROLFILE_LIST",sqltext)):
        document.add_paragraph(ur"(四).数据库物理结构",style='New Heading2')
    elif (re.search("TABLESPACE FRAGMENT",sqltext)):
        document.add_paragraph(ur"(五).表空间",style='New Heading2')
    elif (re.search("CHAINED TABLES",sqltext)):
        document.add_paragraph(ur"(六).表/索引/约束/触发器",style='New Heading2')
    elif (re.search("USER INFO",sqltext)):
        document.add_paragraph(ur"(七).用户定义",style='New Heading2')
    elif (re.search("DB STATS AUTO GATHER",sqltext)):
        document.add_paragraph(ur"(八).统计信息自动收集",style='New Heading2')
    elif (re.search("DB  Block  Buffer  Hit Ratio",sqltext)):
        document.add_paragraph(ur"(九).性能",style='New Heading2')
    # document.add_paragraph(sqltext,"List Bullet 2")
#write sql fetch text to table.
    document.add_paragraph(sqltext,style='New Heading4')
    table = document.add_table(0, len(allrow[0]))
    table.style='Table Grid'
    document.add_paragraph("\n")
    k=k+1
    for row in allrow:
        cells=table.add_row().cells
        col_no = 0
        for col in row:
            try:
                cells[col_no].text=col
                col_no += 1
            except Exception as reason:
                logging.error(reason)
                print("error add table col")
                logging.error("error add table col")

    #set backcolor of first cell begin, 19c need nore. change to 20.
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_3 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_4 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_5 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_6 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_7 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_8 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_9 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_a = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_b = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_c = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_d = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_e = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_f = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_g = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_h = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_i = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_j = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_elm_k = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
    shading_list=[shading_elm_1,shading_elm_2,shading_elm_3,shading_elm_4,shading_elm_5,shading_elm_6,shading_elm_7,shading_elm_8,shading_elm_9,shading_elm_a,shading_elm_b,shading_elm_c,shading_elm_d,shading_elm_e,shading_elm_f,shading_elm_g,shading_elm_h,shading_elm_i,shading_elm_j,shading_elm_k]
    col_cnt=col_no
    col_no=0
    logging.info("column count:"+str(col_cnt))
    while col_no<col_cnt:
        table.rows[0].cells[col_no]._tc.get_or_add_tcPr().append(shading_list[col_no])
        col_no += 1
    #set backcolor of first cell end
    if (re.search("MAX OBJECT ID",sqltext)):
        bitcointable=[[ur"比特币病毒信息"]]
        ps = subprocess.Popen(['grep', 'YES','bitcoin300.out'],stdout=subprocess.PIPE)
        output = ps.communicate()[0]
        for line in output.splitlines():
            bitcointable.append([line])
        if (len(bitcointable) > 1):
            document.add_paragraph("1.1 比特币病毒!!!",style='New Heading4')
            table = document.add_table(0, len(bitcointable[0]))
            table.style='Table Grid'
            document.add_paragraph("\n")
            for row in bitcointable:
                cells=table.add_row().cells
                col_no = 0
                for col in row:
                    try:
                        cells[col_no].text=col
                        col_no += 1
                    except Exception as reason:
                        logging.error(reason)
                        print("error add table bitcoin")
                        logging.error("error add table bitcoin")
        else:
            document.add_paragraph(ur" '1.1 未发现比特币病毒.",style='New Heading4')
        document.add_paragraph(ur"",style='Normal')
        document.add_paragraph(ur"",style='Normal')

        document.add_paragraph(ur" '1.2 SCN信息.",style='New Heading4')
        scntable=[[ur"SCN信息"]]
        ps = subprocess.Popen(['cat', '/tmp/scn.txt'],stdout=subprocess.PIPE)
        output = ps.communicate()[0]
        for line in output.splitlines():
            if not(re.match("--",line)): 
                scntable.append([line])
        if (len(scntable) > 1):
            table = document.add_table(0, len(scntable[0]))
            table.style='Table Grid'
            document.add_paragraph("\n")
            for row in scntable:
                cells=table.add_row().cells
                col_no = 0
                for col in row:
                    try:
                        cells[col_no].text=col
                        col_no += 1
                    except Exception as reason:
                        logging.error(reason)
                        print("error add SCN info")
                        logging.error("error add SCN info")

            #set backcolor of first cell begin
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
            table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_list[col_no])
            #set backcolor of first cell end



    elif (re.search("DB VERSION",sqltext)):

        if (len(opatchtable) > 1):
            document.add_paragraph("Opatch information",style='New Heading5')
            table = document.add_table(0, len(opatchtable[0]))
            table.style='Table Grid'
            document.add_paragraph("\n")
            for row in opatchtable:
                cells=table.add_row().cells
                col_no = 0
                for col in row:
                    try:
                        cells[col_no].text=col
                        col_no += 1
                    except Exception as reason:
                        logging.error(reason)
                        print("error add db version")
                        logging.error("error add db info")

            #set backcolor of first cell begin
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
            table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_list[col_no])
            #set backcolor of first cell end



        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'没有遭遇数据库管理软件bug,[这个版本是oracle 11gR2最终版本；Oracle对此版本的服务截止2020.12月,已经停止.]')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'正常')
    elif (re.search("DB COMP STATUS",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库组件状态都应该是VALID')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'正常')
    elif (re.search("NON-DEFAULT DB PARAMETERS",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库参数设置不能为默认值，需要由资深dba进行具体优化处理')
        document.add_paragraph(ur'此项没有标准得衡量标准，特定应用特定对待，能够满足应用需要得即为合理')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'正常')
    elif (re.search("DB RESOURCE LIMITS",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'MAX_UTILIZATION对应的值应该小于LIMIT_VALUE的值。')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'正常')
    elif (re.search("CONTROLFILE_SIZE",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'控制文件采取2份以上的镜像。')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'正常')
    elif (re.search("ARCHIVELOG RUSH HOUR",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'高峰时没有等待日志切换。')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'正常')
    elif (re.search("PASSWORD_LIFE_TIME",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'用户密码不存在180天有效期限制')
        document.add_paragraph(ur'建议',style='New Heading5')
        if (numrows>0):
            document.add_paragraph(ur'检查用户密码是否将要过期,避免登录失败')
        if (numrows == 0):	
            document.add_paragraph(ur'正常')
    elif (re.search("CREATE SESSION AUDIT",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'如开启审计，不审计普通session 连接。')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'关闭普通连接审计选项')
        document.add_paragraph(ur'正常')
    elif (re.search("AUDIT TRAIL",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'关闭审计,或者明确审计日志存放位置并定期清理')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'如果是db,检查system表空间的v$aud表,如果是OS，经常检查审计日志目录:df-h,df -i')
        document.add_paragraph(ur'正常')
    elif (re.search("RECYCLEBIN",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'回收站不能有太多未清空的表或索引等')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'定期purge recyclebin;清理回收站')
        document.add_paragraph(ur'正常')
    elif (re.search("LOGFILE",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'对没有做RAID的日志文件,每组有2个放在不同磁盘的成员')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph((ur'非冗余系统,建议2个或3个成员.'))
        document.add_paragraph(ur'正常')
    elif (re.search("DATAFILE",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库文件的状态都为ONLINE.')
        document.add_paragraph(ur'手动扩展表空间，不用自动扩展.')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'如果使用自动扩展,监控文件系统空间或者磁盘组空间剩余情况.')
        document.add_paragraph(ur'正常')
    elif (re.search("TEMPFILE",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph( ur'手动扩展临时表空间。临时文件如果自动扩展，易在大排序下消耗存储空间，造成目录填满或性能问题.')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'正常')
    elif (re.search("INVALID DATA FILE",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'系统中不能存在数据文件损坏的情况.')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'正常')
    elif (re.search("TABLESPACE FRAGMENT",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'应该使用本地表空间管理.如果使用字典表空间管理,碎片越少越好.')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'没有返回说明表空间都使用本地管理.')
        document.add_paragraph(ur'正常')
    elif (re.search("DISKGROUP MONITOR",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'磁盘组的空间剩余情况在文件系统中看不见,要经常关注.使用率高于80%时需要准备扩展空间')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph((ur'如果磁盘组使用超过80%,建议扩展'))
        document.add_paragraph(ur'正常')
    elif (re.search("TABLESPACE MONITOR",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'关闭自动扩展的表空间;使用率达到或者超过80％应准备扩展空间')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'正常')
    elif (re.search("UNDO SEGMENT MONITOR",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'回滚段表空间使用率低于80％')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'1.目前参数undo_management=auto 表明数据库处于回滚段自动管理模式下')
        document.add_paragraph(ur'2.只需要监控回滚表空间undo_tablespace=undotbs1的使用率即可')
        document.add_paragraph(ur'3.在回滚表空间使用率接近100％时及时添加数据文件')
        document.add_paragraph(ur'正常')
    elif (re.search("UNDO SEGMENT RATIO",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'回滚空间等待应接近0')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'1目前是自动管理回滚空间,wait应该接近0,否则应扩大回滚表空间')
        document.add_paragraph(ur'正常')
    elif (re.search("BIG SEGMENT",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库中比较大的segment要关注其增长原因和速度')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'关注其增长情况')
        document.add_paragraph(ur'正常')
    elif (re.search("LARGE UNPARTITIONED TABS",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'上述超过2G的表，数据有明显冷热特点，建议对表分区.')
        document.add_paragraph(ur'从物理上将大表分成几个小的分区表，但在逻辑上还是一张表，对于应用透明。')
        document.add_paragraph(ur'这样做的好处是： 1，性能的提高，可以控制数据访问的粒度； 2，数据库可用性提高')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'建议跟应用开发沟通,分区大表')
        document.add_paragraph(ur'正常')
    elif (re.search("DB STATS AUTOTASK WINDOW",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'确认统计信息是按照应用需要,自动或者手动收集')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph((ur'根据应用对数据修改情况,定时主动/或者自动收集统计信息'))
        document.add_paragraph(ur'正常')
    elif (re.search("AWR SNAPSHOT KEEP TIME",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'awr快照频率和保存时间,建议一小时一次,保留30天')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'如果保留时间少于30天,建议扩展到30天.')
        document.add_paragraph(ur'正常')
    elif (re.search("CHAINED TABLES",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库中不存在行迁移或者行链接的表')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'如果有,建议导出导入或者移动表以消除行迁移或者行链接')
        document.add_paragraph(ur'正常')
    elif (re.search("NDEX LEVEL > 3",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库中不索引level大于3')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'如果有,建议重建索引')
        document.add_paragraph(ur'正常')
    elif (re.search("UNUSABLE INDEXES",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库中不存在失效索引')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'如果有,建议重建索引')
        document.add_paragraph(ur'正常')
    elif (re.search("TABLE,INDEX IN SAME TABLESPACE",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库中的表的数据和索引的数据应存放在各自专属表空间')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'如果有,建议将表和索引分开,便于日常维护')
        document.add_paragraph(ur'正常')
    elif (re.search("UNSYSTEM OBJ IN SYSTEM",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'系统表空间中不存放任何用户数据')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'建议将这些数据移动到用户表空间')
        document.add_paragraph(ur'正常')
    elif (re.search("INVALID CONSTRAINTS",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库中不存在失效约束')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'跟应用确认,如果需要,重建或者这些约束,如果不需要可以删除.')
        document.add_paragraph(ur'正常')
    elif (re.search("INVALID TRIGGERS",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库中不存在失效触发器')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'跟应用确认,如果需要,重建或者这些触发器,如果不需要可以删除.')
        document.add_paragraph(ur'正常')
    elif (re.search("INVALID OBJECTS",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库中不存在失效对象')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'跟应用确认,如果需要,重建或者这些对象,如果不需要可以删除.')
        document.add_paragraph(ur'正常')
    elif (re.search("USER INFO",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'数据库中不使用的用户状态为LOCK')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'对于上述开启的帐户确认是否应用需要的帐户，对于不使用的帐户建议lock或删除')
        document.add_paragraph(ur'正常')
    elif (re.search("SUPER USERS",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'拥有数据库启停权限的用户只能是sys')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'上述用户拥有较高的管理角色权限，注意此类用户的密码控制')
        document.add_paragraph(ur'正常')
    elif (re.search("DBA PRIVS",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'dba权限.拥有DBA权限的用户,是按照需求授予的')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'没有非需求用户被授予dba权限')
        document.add_paragraph(ur'正常')
    elif (re.search("SYS PRIVS",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'sys权限.拥有SYS权限的用户,是按照需求授予的')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'没有非需求用户被授予dba权限')
        document.add_paragraph(ur'正常')
    elif (re.search("OBJECT PRIVS",sqltext)):
        document.add_paragraph(ur'健康指标',style='New Heading5')
        document.add_paragraph(ur'对象权限.数据库中的对象权限是按照需求授予的')
        document.add_paragraph(ur'建议',style='New Heading5')
        document.add_paragraph(ur'如果存在不合理数据访问，需回收相应对象权限(限于文档大小,只列10条)')
        document.add_paragraph(ur'正常')

document.add_page_break()
ostitle=ur"(九). AWR报告"
document.add_paragraph(ostitle ,style='New Heading2')
document.add_paragraph('')
document.add_paragraph('')
document.add_paragraph(ur'系统性能良好')

document.add_page_break()
ostitle=ur"(十). 告警日志"
document.add_paragraph(ostitle ,style='New Heading2')
try:
    ps = subprocess.Popen(['tail','-n80','dberr.txt'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]
except Exception as reason:
    logging.error(reason)
    print("cat  dberr.txt fail")
    logging.error("error add db alert")
for line in output.splitlines():
    try:
        document.add_paragraph(line,style='linuxoutput')
    except Exception as reason:
        logging.error(reason)
        print("add dberr fail")
        logging.error("error add db err paragraph  ")
document.add_paragraph(ur'近期未见严重报错')
#*******************This part get OS information ***************************
#os information
document.add_page_break()

#os information
try:
    ps = subprocess.Popen(['hostname'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]
except Exception as reason:
    logging.error(reason)
    print("hostname fail")
    logging.error('get hostname fail.')

for line in output.splitlines():
    print line;

hostname=line;
filename=filename+"_"+dbnamestr+"_"+hostname+"_"+filename_end
ostitle=ur"五. "+hostname+ "  硬件信息"
document.add_paragraph(ostitle ,style='New Heading1')


osver=ur"(一). OS版本"
document.add_paragraph(osver ,style='New Heading2')

try:
    ps = subprocess.Popen(['cat', '/etc/redhat-release'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]
except Exception as reason:
    logging.error(reason)
    print("cat redhat release fail")
    logging.error('get redhat-releae fail.')

for line in output.splitlines():
    document.add_paragraph(line,style='linuxoutput')

try:
    ps = subprocess.Popen(['uname','-r'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]
except Exception  as reason:
    logging.error(reason)
    print("uname -r fail")
    logging.error('uname -r fail.')
for line in output.splitlines():
    document.add_paragraph(line,style='linuxoutput')


meminfo=ur"(二). 内存"
document.add_paragraph(meminfo ,style='New Heading2')
try:
    ps = subprocess.Popen(['free','-h'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]
except Exception as reason:
    logging.error(reason)
    print("free -h fail")
    logging.error('free -h fail.')
for line in output.splitlines():
    document.add_paragraph(line,style='linuxoutput')

meminfo=ur"(三). 磁盘IO(iostat) "
document.add_paragraph(meminfo ,style='New Heading2')
#document.add_paragraph((meminfo,style='List Bullet'))
try:
    ps = subprocess.Popen(['iostat','2','2'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]
except Exception as reason:
    logging.error(reason)
    print("iostat fail")
    logging.error('iostat  fail.')
for line in output.splitlines():
    try:
        document.add_paragraph(line,style='linuxoutput')
    except Exception as reason:
        logging.error(reason)
        print("iostat fail")
        logging.error('iostat  fail.')


#output top cpu usage

ostitle=ur"六. "+hostname+ "  OS性能"
document.add_paragraph(ostitle ,style='New Heading1')
meminfo=ur"(一). CPU当前LOAD"
document.add_paragraph(meminfo ,style='New Heading2')
try:
    p1 = subprocess.Popen(['top', 'b', '-n1'],stdout=subprocess.PIPE)
    p2 = subprocess.Popen(['head', '-5'],stdin=p1.stdout,stdout=subprocess.PIPE)
    output = p2.communicate()[0]
except Exception as reason:
    logging.error(reason)
    print("top fail")
    logging.error('top  fail.')
for line in output.splitlines():
    document.add_paragraph(line,style='linuxoutput')

meminfo=ur"(二). CPU历史 LOAD"
document.add_paragraph(meminfo ,style='New Heading2')
meminfo=ur"过去30天CPU LOAD前20,取自dba_hist_osstat"
document.add_paragraph(meminfo)

sql="select to_char(round(s.end_interval_time, 'hh24'), 'yyyy-mm-dd hh24') snap_time, os.instance_number,os.value \"CPU_LOAD(%)\" from dba_hist_snapshot s, dba_hist_osstat os where s.dbid = os.dbid and s.instance_number = os.instance_number and s.snap_id = os.snap_id and os.stat_name = 'LOAD' AND S.END_INTERVAL_TIME between sysdate-30 and sysdate and rownum <= 10 order by os.value desc ,to_char(trunc(s.end_interval_time, 'hh24'), 'yyyy-mm-dd hh24'), os.instance_number"
cur = con.cursor()
try:
    cur.execute(sql)
except Exception  as reason:
      logging.error(sql)
      print sql
      print "Execute error ."  
      onerow="    "
try:
    result = cur.fetchall()
except Exception  as reason:
    logging.error(reason)
    print "execute fail"
    logging.error(sql)
col_names = [row[0] for row in cur.description or []]
numcols=len(col_names)
numrows=len(result)
allrow=[col_names]
i=0
while i< numrows:
    j=0
    newrow=[]
    while j<numcols:
        col=str(result[i][j])
        newrow.append(col)
        j=j+1
    allrow.append(newrow)
    newrow=[]
    i=i+1

table = document.add_table(0, len(allrow[0]))
table.style='Table Grid'
for row in allrow:
    cells=table.add_row().cells
    col_no = 0
    for col in row:
        cells[col_no].text=col
        col_no += 1
#set backcolor of first cell begin
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_2 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_elm_3 = parse_xml(r'<w:shd {} w:fill="99CCFF"/>'.format(nsdecls('w')))
shading_list=[shading_elm_1,shading_elm_2,shading_elm_3]
col_no=0
while col_no<3:
    table.rows[0].cells[col_no]._tc.get_or_add_tcPr().append(shading_list[col_no])
    col_no += 1
#set backcolor of first cell end

#test  MEM usage top 10
#"ps aux | awk '{print $2, $4, $6, $11}' | sort -k3rn | head -n 10"
document.add_paragraph((""))

meminfo=ur"(三). 内存占用top 10进程"
document.add_paragraph(meminfo ,style='New Heading2')
try:
    p1 = subprocess.Popen(['ps','aux'],stdout=subprocess.PIPE)
    p2 = subprocess.Popen(['awk', '"{print $2, $4, $6, $11}"' ],stdin=p1.stdout,stdout=subprocess.PIPE)
    p3 = subprocess.Popen(['sort', '-k3rn' ],stdin=p2.stdout,stdout=subprocess.PIPE)
    p4 = subprocess.Popen(["head", "-n","10" ],stdin=p3.stdout,stdout=subprocess.PIPE)
    output = p4.communicate()[0]
except Exception  as reason:
    logging.error(reason)
    print("ps fail")
    logging.error("ps fail")
for line in output.splitlines():
    document.add_paragraph(line,style='linuxoutput')

#test CPU usage top 10
#ps = subprocess.Popen(["top b -n1 | head -17 | tail -11"],stdout=subprocess.PIPE)
meminfo=ur"(四). CPU 占用top 10进程"
document.add_paragraph(meminfo ,style='New Heading2')
try:
    p1 = subprocess.Popen(['top','b','-n1'],stdout=subprocess.PIPE)
    p2 = subprocess.Popen(['head','-17'],stdin=p1.stdout,stdout=subprocess.PIPE)
    p3 = subprocess.Popen(['tail', '-11' ],stdin=p2.stdout,stdout=subprocess.PIPE)
    output = p3.communicate()[0]
except Exception as reason:
    logging.error(reason)
    print("top fail")
    logging.error("top fail")
for line in output.splitlines():
    document.add_paragraph(line,style='linuxoutput')

# OS health check
document.add_paragraph(ur'健康指标',style='New Heading5')
document.add_paragraph(ur'CPU使用率（LOAD）要长时间低于70％; 内存不出现使用虚拟内存的现象;  IO WAIT 小于15％')
document.add_paragraph(ur'建议',style='New Heading5')
document.add_paragraph(ur'正常')


osdk=ur"(五). 磁盘使用率(df -Ph)"
document.add_paragraph(osdk ,style='New Heading2')

try:
    ps = subprocess.Popen(['df','-Ph'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]
except Exception as reason:
    logging.error(reason)
    print("df -Ph fail")
    logging.error("df -Ph fail")
for line in output.splitlines():
    try:
        document.add_paragraph(line,style='linuxoutput')
    except Exception as reason:
        logging.error(reason)
        print("df -Ph fail")
        logging.error("df -Ph fail")

osdk=ur"(六). inode使用情况df -Pi"
document.add_paragraph(osdk ,style='New Heading2')
try:
    ps = subprocess.Popen(['df','-Pi'],stdout=subprocess.PIPE)
    output = ps.communicate()[0]
except Exception as reason:
    logging.error(reason)
    print("df -Pi fail")
    logging.error("df -Pi fail")
for line in output.splitlines():
    try:
        document.add_paragraph(line,style='linuxoutput')
    except Exception as reason:
        logging.error(reason)
        print("df -Pi fail")
        logging.error("df -Pi fail")

# OS health check
document.add_paragraph(ur'健康指标',style='New Heading5')
document.add_paragraph(ur'数据库安装软件目录和数据文件存放目录使用率不超过70%;   每个目录inode使用率不超过70%')
document.add_paragraph(ur'建议',style='New Heading5')
document.add_paragraph(ur'正常')

document.add_page_break()
ostitle=ur"七. 总结和建议"
document.add_paragraph(ostitle ,style='New Heading1')

ostitle=ur"(一). 巡检总结"
document.add_paragraph(ostitle ,style='New Heading2')

sumall=ur"数据安全性:"
document.add_paragraph(sumall,style='New Heading5')
document.add_paragraph((ur"*. 部署RMAN备份: "))
document.add_paragraph((ur"*. 部署数据泵导出: "))
document.add_paragraph((ur"*. 部署同步到灾备: "))
document.add_paragraph((ur"*. 部署dataguard: "))
sumall=ur"数据库稳定性:"
document.add_paragraph(sumall,style='New Heading5')
document.add_paragraph((ur"数据库版本稳定. oracle 已经停止对11g支持。建议测试Oracle 19c."))
sumall=ur"数据库可用性:"
document.add_paragraph(sumall,style='New Heading5')
document.add_paragraph((ur"*. 部署RAC: "))
document.add_paragraph((ur"*. 部署同步到灾备: "))
document.add_paragraph((ur"*. 部署dataguard: "))

sumall=ur"数据库性能:"
document.add_paragraph(sumall,style='New Heading5')
document.add_paragraph(ur"数据库目前性能良好.")

ostitle=ur"(二). 主要建议"
document.add_paragraph(ostitle ,style='New Heading2')
document.add_paragraph(ur"数据安全性. ",style='New Heading5')

document.add_paragraph(ur"*. 检查rman备份日志，确保备份成功，归档空间被清理；")
document.add_paragraph(ur"*. 定期rman恢复测试，确保rman备份可用.")
document.add_paragraph(ur"*. 定期数据泵导入测试，确保备份可用.")
document.add_paragraph(ur"*. 定期DataGuard切换测试.")
document.add_paragraph(ur"*. 定期灾备库切换，确保它们在需要时可用.")

document.add_paragraph(ur"数据库可用性.", style='New Heading5')
document.add_paragraph(ur"数据库性能.",style='New Heading5')

################
document.save(filename+'.docx')
print "done."

